import pandas as pd

# read the following excel file "C:\Users\Brooks Scarpa\Downloads\Word- GENERAL_NOTES_RESI_2020 CITY OF CulverC.xlsx" into a dataframe

import openpyxl
from docx import Document
from docx.shared import Inches



def convert_excel_to_word(excel_file, word_file):
    # Load the Excel workbook
    wb = openpyxl.load_workbook(excel_file)
    ws = wb.active

    # Create a new Word document
    doc = Document()
    
    # Loop through rows in Excel worksheet
    for row in ws.iter_rows(min_row=1, values_only=True):
        # Calculate the level of indentation based on the row index
        indentation_level = row[0].count('.') if row[0] else 0
        
        # Create a new paragraph in the Word document
        para = doc.add_paragraph()

        # Apply indentation to the paragraph
        para.paragraph_format.left_indent = indentation_level * 1  # Adjust the multiplier as needed
        
        # Add cell values to the paragraph
        for cell_value in row:
            para.add_run(str(cell_value) + ' ')
    
    # Save the Word document
    doc.save(word_file)
    
# Replace 'input.xlsx' with your Excel file name and 'output.docx' with your desired Word file name

def clean_excel(path,name):
    
    df = pd.read_excel(path)

    #remove all empty rows from the dataframe
    df = df.dropna(how='all')

    #remove all empty columns from the dataframe
    df = df.dropna(axis=1, how='all')
    
    file_path = name + '.txt'
    
    with open(file_path, 'w') as f:
        
        for index, row in df.iterrows():
            
            #concatenate the values in the cell to a string, each empty cell shall be replave with a tabulated space
            row = row.fillna('\t')
            row = row.astype(str)
            row = row.str.cat(sep='')
            
            print(row)

            #write the string to a text file

            f.write(row+'\n')
            

if __name__ == '__main__':
    
    #file_name = '/2022 CBC GREEN BLDG CODE - NONRESI MANDATORY MEASURES.xlsx'
    #file_name = '/2022 CBC GREEN BLDG CODE - RESI MANDATORY MEASURES.xlsx'
    file_name = '/2022 CBC_B+S Resi, Commercial & Mixed Use Accessibility Notes_CH11B.xlsx'
    folder = 'C:/Users/Brooks Scarpa/Documents/GitHub/brooks-scarpa-repo/General_Notes'
    
    name = file_name.split('.')[0]
    name = folder + '/' + name
    
    #clean_excel(folder+file_name,name)
    
    convert_excel_to_word(folder+file_name, name + '.docx')